from docx import *
import re

emails = []
rows = []
document = Document('parentList.docx')
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text == "Mother":
                    rows.append(row)

for row in rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if not paragraph.text == "":
                email_list = re.findall(r"[a-z0-9\-+_]+@[a-z0-9\-+_]+\.[a-z]+", paragraph.text)
                for email in email_list:
                    emails.append(email)

output = Document()
for email in emails:
    p = output.add_paragraph(email)

output.save('motherEmailList.docx')
