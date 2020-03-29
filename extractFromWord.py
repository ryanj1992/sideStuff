from docx import *
import re
import json
from pprint import pprint

document = Document('randomNames.docx')


class Person:
    def __init__(self, fullname, fullemail, fullparent):
        self.name = fullname
        self.email = fullemail
        self.parent = fullparent


objs = list()
for para in document.paragraphs:
    text = para.text
    if "Mother" in text:
        name_list = re.findall(r"^([\w]+ [\w]+)", text)
        email_list = re.findall(r"[a-z0-9\-+_]+@[a-z0-9\-+_]+\.[a-z]+", text)
        parent_list = re.findall(r"^(?:\S+\s){2}(\S+)", text)
        for name in name_list:
            for email in email_list:
                for parent in parent_list:
                    objs.append(Person(name, email, parent))

# for obj in objs:
#     print(obj.name + " " + obj.email + " " + obj.parent)

output = Document()

table = output.add_table(len(objs), cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Email'
hdr_cells[2].text = 'Parent'

for obj in objs:
    row_cells = table.add_row().cells
    row_cells[0].text = obj.name
    row_cells[1].text = obj.email
    row_cells[2].text = obj.parent

output.save('parentList.docx')

# output = Document()
# for name in names:
#     for parent in parents:
#         for email in emails:
#             p = output.add_paragraph(name + parent + email)
# output.save('emails.docx')








